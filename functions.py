import json
import docx
from typing import List, Dict, Optional, Tuple
import pandas as pd
import numpy as np
from datetime import datetime
from copy import deepcopy
import os
import shutil


def load_json(path: str) -> List[Dict[str, str]]:
    """
    Load JSON file, return data as list of dicts.
    """
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def save_json(path: str,
              data: List[Dict[str, str]]) -> None:
    """
    Save list of dicts to JSON file.
    """
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)


def get_doc_text(path: str) -> List[str]:
    """
    Extract text from document at path.
    """
    return [['','[B]'][any([j in p.paragraph_format.element.xml
                            for j in ['<w:numPr>', 'w:val="ListBullet"']])]
            + p.text
            for p in docx.Document(path).paragraphs]


def get_line_index(findtxt: str, 
                   intxt: str) -> int:
    """
    Return line number within findtxt that begins with intxt.
    """
    return (z[0] if (z:=[i for i,j in enumerate(intxt) 
                         if j.upper().startswith(findtxt.upper())]) 
            else -1)


def normalize_text(text: str) -> str:
    """
    Normalize text by replacing special characters and trimming whitespace.
    """
    if text is None:
        return ""
    repl_map = [("\u201c", '"'), 
                ("\u201d", '"'), 
                ("\u2018", "'"), 
                ("\u2019", "'"), 
                ("\u2013", "-"), 
                ("\u2014", "-")]
    return [(z:=(z if i else text).replace(j,k)) 
            for i,(j,k) in enumerate(repl_map)][-1].strip()


def past_example_components(past_examples: List[Dict[str, str]]) -> List[Dict[str, str]]:
    """
    Get details of past applications from JSON file, read 
    """
    doc_components = []
    for e in past_examples:
        resume_text = get_doc_text(e['resume_path'])
        cover_text = get_doc_text(e['cover_path'])
        doc_components += [{
            'r_title': resume_text[2:3],
            'r_hl_skills': resume_text[3:4],
            'r_intro': resume_text[4:get_line_index('Technical Skills', resume_text)-1],
            'r_tech': resume_text[get_line_index('Technical Skills', resume_text)+1:
                                  get_line_index('Areas of Expertise', resume_text)-1],
            'r_expert': resume_text[get_line_index('Areas of Expertise', resume_text)+1:
                                    get_line_index('Professional Experience', resume_text)-1],
            'r_effo': resume_text[get_line_index('Data Science Director', resume_text)+2:
                                  get_line_index('dunnhumbyUSA', resume_text)-1],
            'r_dusa_vp': resume_text[get_line_index('Vice President', resume_text)+2:
                                     get_line_index('Analysis Director', resume_text)-1],
            'r_dusa_dir': resume_text[get_line_index('Analysis Director', resume_text)+2:
                                      get_line_index('Senior Analyst', resume_text)-1],
            'r_dusa_sa': resume_text[get_line_index('Senior Analyst', resume_text)+2:
                                     get_line_index('dunnhumby ', resume_text)-1],
            'r_duk': resume_text[get_line_index('Marketing Analyst', resume_text)+2:
                                 get_line_index('Education', resume_text)-1],
            'c_p1': cover_text[(s:=get_line_index('Dear', cover_text)+1): (s:=s+1)], 
            'c_p2': cover_text[s: (s:=s+1)], 
            'c_p2b': cover_text[s: (s:=s+sum([i[:3]=='[B]' for i in cover_text[s:]]))], 
            'c_p3': cover_text[s: (s:=s+1)], 
            'c_p4': cover_text[s: (s:=s+1)], 
        }]
    doc_components = [{k: [normalize_text(l) for l in e[k] if l] 
                       for k in e} 
                      for e in doc_components]
    fix_bullet_lists(doc_components)
    
    return doc_components


def fix_bullet_lists(doc_components):
    """
    Some far-from-ideal code that identifies instances where some of the bullets in the
    last-produced resume/cover letter were manually updated, and adds the updated text
    to the input/bullets_....txt files.
    """
    for key in ('r_effo', 'r_dusa_vp', 'r_dusa_dir', 'r_dusa_sa', 'r_duk', 'c_p2b'):
        # Read in file with previously-used text for bullet points
        bullets_hist = [[v.strip() for v in g.split('\n')] 
                        for g in open(f'input/bullets_{key}.txt', 'r')
                                     .read().split('\n\n')]
        
        # Remove any that were never used (replaced by manual edits)        
        b_used = {b.replace('[B]','') for e in doc_components for b in e[key]}
        bullets_hist = [[b for b in g if b in b_used] for g in bullets_hist]
        
        # Identify bullets in last-produced document that don't match to bullets_hist
        # (i.e. results of manual edits)
        b_matched = [False for i in bullets_hist]
        edited = []
        for b in doc_components[-1][key]:
            bs = b.replace('[B]','')
            if any(z:=[bs in i for i in bullets_hist]):
                b_matched[z.index(True)] = True
            else:
                edited += [bs]
                
        # Add each manually-edited bullet to the correct group
        # (using quick-and-dirty matching method - % of words matched to group)
        if edited:
            bh_match = [[i, {w.strip('.,').lower()
                             for k in bullets_hist[i] 
                             for w in k.split()}] 
                        for i,j in enumerate(b_matched) if not j]
            for e in edited:
                ix = sorted([[sum(z:=[w.strip('.,').lower() in j 
                                      for w in e.split()])/len(z), 
                              i]
                             for i,j in bh_match])[-1][1]
                bullets_hist[ix] += [e]
                
        # Save updated bullets data
        with open(f'input/bullets_{key}.txt', 'w') as f:
            f.write('\n\n'.join(['\n'.join(g) for g in bullets_hist]))


def get_job_details() -> Dict[str, str]:
    """
    Gather job posting details from user inputs, clean and return as dict.
    """
    title_input = company_input = loc_input = details_input = comments_input = ''
    while not title_input:
        title_input = input('Enter job title')
    while not company_input:
        company_input = input('Enter name of hiring company')
    while not loc_input:
        loc_input = input('Enter location of hiring company')
    while not details_input:
        details_input = input('Enter job description')
        
    comments_input = input('Enter special comments')
    if comments_input:
        comments_input = ("Here are some additional comments, describing points that I think "
                          + "may be worth emphasizing when generating your response\n"
                          + normalize_text(comments_input)
                          + "\n(END OF COMMENTS)\n")
        
    return {
        "role": normalize_text(title_input),
        "company": normalize_text(company_input).strip('.'),
        "location": normalize_text(loc_input),
        "apply_date": datetime.today().strftime("%m%d%Y"),
        "job_posting": normalize_text(details_input),
        "comments": comments_input,
        }


def get_prompts() -> Dict[str, str]:
    """
    Read in prompt templates from .txt files and return in dict
    """
    return {f.replace('.txt',''): open(f'prompts/{f}', 'r').read()
            for f in os.listdir('prompts')}


def proper(text: str) -> str:
    """
    Apply sentence case to string, but preserve acronyms as all-caps.
    """
    f,t = 0,1
    while f<len(text):
        while t<len(text) and (65<=ord(text[t])<=90 or 97<=ord(text[t])<=122):
            t+=1
        if text[f:t]!=text[f:t].upper():
            text = (text[:f]
                    +(text[f:t].lower() if f else text[f:t].capitalize())
                    +text[t:])
        t = (f:=t+1)+1
    return text


def get_source_examples(doc_components: List[Dict[str, str]], 
                        key: str) -> str:
    """
    Extract text for source material examples for component of resume/cover letter
    """
    if key in ('r_title'):
        return '\n'.join({i.upper() 
                          for i in sum([e[key] 
                                        for e in doc_components],[])})
    elif key in ('r_hl_skills', 'r_expert'):
        return '\n'.join({proper(s.strip())
                          for e in doc_components
                          for s in '|'.join(e[key]).split('|')})
    elif key in ('r_tech'):
        exnum, sk, py = zip(*[[n]+[i.replace(')','').strip() 
                                   for i in (s.split('(')+[''])[:2]]
                              for n, e in enumerate(doc_components)
                              for s in '|'.join(e[key]).split('|')])
        skord = [(z:=1 if i==0 or j!=exnum[i-1] else z+1) for i,j in enumerate(exnum)]
        source_examples = '\n'.join(pd.DataFrame({'skill': sk, 'skrank': skord})
                                    .groupby('skill').mean()
                                    .sort_values(by='skrank').index.tolist())
        source_examples += '\n\nPYTHON PACKAGES:\n'
        source_examples += '\n'.join(pd.DataFrame([[k.strip(),j]
                                                   for i in py if i
                                                   for j,k in enumerate(i.split(','))],
                                                  columns=['pyskill', 'skrank'])
                                     .groupby('pyskill').mean()
                                     .sort_values(by='skrank').index.tolist())
        return source_examples
    elif key in ('r_effo', 'r_dusa_vp', 'r_dusa_dir', 'r_dusa_sa', 'r_duk', 'c_p2b'):
        return '\n-----\n'.join([l 
                                 for l in open(f'input/bullets_{key}.txt', 'r')
                                              .read().strip().split('\n\n')])
    else:
        return '\n\n'.join([e[key][0] for e in doc_components])


def split_llm_outputs(raw_llm_outputs: Dict[str, List[str]]) -> Tuple[Dict[str, List[str]], str]:
    """
    Separate out feedback on how much new text was composed by AI from llm_outputs dict
    """
    llm_outputs = {}
    llm_report = {}
    for n,key in enumerate(raw_llm_outputs):
        if key in ['r_effo', 'r_dusa_vp', 'r_dusa_dir', 'r_dusa_sa', 'r_duk', 'c_p2b']:
            gentxt = [([j.strip() for j in i.split('\n')]+[*'XX'])[:3]
                      for i in raw_llm_outputs[key].split('\n\n')]
            brpt, btxt, brnk = [*zip(*gentxt)]
            
            # bullets_hist = [[v.strip() for v in g.split('\n')] 
            #                 for g in open(f'input/bullets_{key}.txt', 'r')
            #                               .read().split('\n\n')]
            # bh_chg = False
            # for bn in range(len(bullets_hist)):
            #     if btxt[bn] != 'X' and btxt[bn] not in bullets_hist[bn]:
            #         bullets_hist[bn] += [btxt[bn]]
            #         bh_chg = True
            # if bh_chg:
            #     with open(f'input/bullets_{key}.txt', 'w') as f:
            #         f.write('\n\n'.join(['\n'.join(g) for g in bullets_hist]))
                    
            llm_report[key] = ''.join([f"{' '*8}Bullet {i}: {j}\n" 
                                       for i,j in sorted(zip(brnk, brpt)) 
                                       if i != 'X'])
            llm_outputs[key] = [j 
                                for i,j in sorted(zip(brnk, btxt)) 
                                if i != 'X']
        else:
            fb, *llm_outputs[key] = raw_llm_outputs[key].split('\n')
            llm_report[key] = f"{' '*8}{fb}\n"
    llm_report_txt = '\n'.join([f"{k}\n{v}" for k,v in llm_report.items()])
            
    return llm_outputs, llm_report_txt


def make_doc(label: str, 
             job_details: Dict[str, str], 
             llm_outputs: Dict[str, List[str]]) -> docx.document.Document:
    """
    Populate customized text into document template
    """
    template = docx.Document(f'templates/{label}_template.docx')

    doc = deepcopy(template)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    
    for p in template.paragraphs:
        if p.text.startswith('<'):
            for l in llm_outputs.get(p.text[1:-1], 
                                     [job_details.get(p.text[1:-1])]):
                doc.add_paragraph(l, p.style)
        else:
            doc.add_paragraph(p.text, p.style)
        if "w:type=\"page\"" in p._element.xml:
            doc.add_page_break()
            
    return doc


def save_to_folder(job_details: Dict[str, str], 
                   resume: docx.document.Document, 
                   coverletter: docx.document.Document, 
                   llm_report: str) -> List[Dict[str, str]]:
    """
    Write generated resume and cover letter to Documents/Job Search folder, with AI report.
    Also return new entry at add to past examples json.
    """
    savepath = os.path.join("/mnt/c/Users/gavin/OneDrive/Documents/Job Search",
                            job_details['company'])
    if not os.path.exists(savepath):
        os.makedirs(savepath)
        
    tmp_path = os.path.join(savepath, 'temp.docx')
    
    resume.save(tmp_path)
    resume_path = os.path.join(savepath, 
                               f"Gilchrist_Gavin_Resume_{job_details['apply_date']}.docx")
    shutil.move(tmp_path, resume_path)
    
    coverletter.save(tmp_path)
    coverletter_path = os.path.join(savepath, 
                                    f"Gilchrist_Gavin_CoverLetter_{job_details['apply_date']}.docx")
    shutil.move(tmp_path, coverletter_path)
    
    with open(os.path.join(savepath, 
                           f"LLM_Report_{job_details['apply_date']}.txt"), 'w') as f:
        f.write(llm_report)
        
    return [{**job_details,
             "resume_path": resume_path,
             "cover_path": coverletter_path}]
